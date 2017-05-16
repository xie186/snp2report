# -*- coding: utf-8 -*-
import re
from pubmed import PubMed
from metapub import PubMedFetcher
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

class ReportDoc():
    """docstring for ClassName"""
    def __init__(self):
        pass

    def readTemplate(self, options):
        """
        Read template 
        """
        #document = Document()
        document = Document(options.template)
        return document
 
    def generRep(self, snp_data, var_type, var_description, document):
        """
        
        """ 
        class_var = {}
        for snp in snp_data.iterrows():
            #["Gene", "RSID", "Study", "Genotype", "Judgement"])
            rsid = snp[1]['RSID']
            geno = snp[1]['Genotype']
            if rsid not in var_type:
                print rsid + "\t" + geno
            else:
                if rsid not in var_type:
                    print rsid + "not in var_type"
                class_str = var_type[rsid]
                if class_str not in class_var:
                    class_var[class_str] = []
                class_var[class_str].append(snp[1])
        for class_str in class_var:
            p = document.add_paragraph()
            header = document.add_heading(class_str, level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for snp in class_var[class_str]: 
                regexp = re.compile(ur'正常|平均') 
                if not regexp.search(snp['Judgement']):
                    document.add_heading(snp['Study'], level=2)
                    text = snp['Study'] +  snp['Judgement']
                    self.processPMID(var_description[snp['RSID']], document, text)
                    #document.add_paragraph(var_description[snp['RSID']])
                     

    def processPMID(self, description, document, text):
        """XXX"""
        
        pmid = re.compile('PMID *(\d+)') 
        list_pmid = pmid.findall(description)
        description = re.sub(r'\[PMID *\d+\]', '', description)
        pmcid = re.compile('PMCID *(\d+)')
        list_pmcid = pmcid.findall(description)
        description = re.sub(r'\[PMCID *\d+\]', '', description)
        para = description.split(ur'\n')
        for para_str in para:
            #print para_str
            p = document.add_paragraph('	')
            p.add_run(para_str)
        std_str = u"我们通过检测您的基因位点，使用PUBMED等国际公认参考系统，我们认为" + text + u"。"
        p = document.add_paragraph('	')
        p.add_run(std_str)
      
        fetch = PubMedFetcher()
        for pmid in list_pmid:
            # http://www.ncbi.nlm.nih.gov/pubmed/26471457
            pm = fetch.article_by_pmid(pmid)
            title = pm.title
            title = re.sub('\.', '', title)
            citation = '. '.join([title, pm.journal])
            p = document.add_paragraph()
            p.add_run(citation).italic = True
       
        for pmcid in list_pmcid:
            pm = fetch.article_by_pmcid(pmcid)
            title = pm.title
            title = re.sub('\.', '', title)
            citation = '. '.join([title, pm.journal])
            p = document.add_paragraph()
            p.add_run(citation).italic = True

    def generTable(self, snp_data, snp_geno_flag, document):
        """
        Input: snp_data: pandas.core.frame.DataFrame
               snp_geno_flag: snp_geno_flag
        """
        snp_sig = {}
        p = document.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        table_header = u'基因检测' + str(len(snp_data)) + u'项数据汇总'
        #table_header.decode('utf-8')

        header = document.add_heading(table_header, level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(rows= 1, cols=4, style=None)
        table.style = 'Table Grid' #Table Grid
        print len(snp_data)
        #print snp_data.head()
        hdr_cells = table.rows[0].cells
        text_list = [u'基因', u'SNP位点', u'基因型', u'影响']
        for i in range(0, len(text_list)):
            self.outputHeaderCell(hdr_cells, i, text_list[i])

        #hdr_cells[0].paragraphs[0].add_run('')
        styles = document.styles
        #style = styles.add_style('Light Grid Accent 1', WD_STYLE_TYPE.TABLE)
        table_styles = [
                        s for s in styles if s.type == WD_STYLE_TYPE.TABLE
                ]
        for style in table_styles:
            pass
            #print(style.name)
        for snp in snp_data.iterrows():
            row_cells = table.add_row().cells
            #["Gene", "RSID", "Study", "Genotype", "Judgement"])
            color = self.getColor(snp[1])
            text = snp[1]['Study'] +  snp[1]['Judgement']
            snp_sig[snp[1]['RSID']] = text
            run = row_cells[0].paragraphs[0].add_run(snp[1]['Gene'])
            font = run.font
            font.color.rgb = color
            run = row_cells[1].paragraphs[0].add_run(snp[1]['RSID'])
            font = run.font
            font.color.rgb = color
            run = row_cells[2].paragraphs[0].add_run(snp[1]['Genotype'])
            font = run.font
            font.color.rgb = color
            run = row_cells[3].paragraphs[0].add_run(text)
            font = run.font
            font.color.rgb = color
        table.style = 'Table Grid'

    def getColor(self, snp):
        color = None
        regexp = re.compile(ur'用药谨慎|风险偏高|需求增加|影响较大')
        if regexp.search(snp['Judgement']):
            color = RGBColor(255, 69, 0) 
        regexp = re.compile(ur'可能较好|概率偏高')
        if regexp.search(snp['Judgement']):
            color = RGBColor(100, 149, 237)
        return color
        
    def outputHeaderCell(self, hdr_cells, i, text):
       """Output header cell"""
       run = hdr_cells[i].paragraphs[0].add_run(text)
       font = run.font
       font.bold = True
       #font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

