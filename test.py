# -*- coding: utf-8 -*-

#import urllib3.request
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

document = Document()

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

#urllib3.request.urlretrieve("http://placehold.it/350x150", "placeholder.png")
document.add_picture('monty-truth.png', width=Inches(1.25))

recordset = [
    {
        "id" : 1,
        "qty": 2,
        "desc": "New item"
    },
    {
        "id" : 2,
        "qty": 2,
        "desc": "New item"
    },
    {
        "id" : 3,
        "qty": 2,
        "desc": "New item"
    },

]

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells

hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
#hdr_cells[2].paragraphs[0].text = 'XX'
#print dir(hdr_cells[2].paragraphs[0])
#para = hdr_cells[2].add_paragraph('')
run = hdr_cells[2].paragraphs[0].add_run('XX')
font = run.font
font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

shading_elm = parse_xml(r'<w:shd {} w:fill="dce3db"/>'.format(nsdecls('w')))
hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)


for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item["qty"])
    row_cells[1].text = str(item["id"])
    row_cells[2].text = item["desc"]

## http://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
table.style = 'TableGrid'
document.add_page_break()
document.save('test.docx')
